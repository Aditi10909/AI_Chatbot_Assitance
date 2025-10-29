import io
from typing import List, Dict, Tuple
import os
from logger import logging
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
import textwrap

logger = logging.getLogger(__name__)


# Optional imports (try to import, else fallback)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from fpdf import FPDF
except Exception:
    FPDF = None

def parse_sections_from_query(query: str, available_sections: List[str]) -> Tuple[List[str], List[str]]:
    """
    simple string matching + splitting by commas/and.
    Returns (found_sections, missing_sections).
    """
    q = query.lower()
    # common separators
    tokens = [t.strip() for t in q.replace(" and ", ",").split(",")]
    tokens = [t for t in tokens if 1 < len(t) < 100]
    found = []
    for s in available_sections:
        sl = s.lower()
        # exact substring match or token match
        if sl in q or any(sl in t for t in tokens) or any(t in sl for t in tokens):
            found.append(s)
    # If parser couldn't find explicit sections, try heuristics: look for words like "introduction", "summary", "table"
    if not found:
        heuristics = ["introduction", "summary", "findings", "clinical", "tables", "images", "results", "discussion"]
        for h in heuristics:
            if h in q:
                # pick available sections that contain that heuristic
                for s in available_sections:
                    if h in s.lower() and s not in found:
                        found.append(s)

    # If still empty, attempt to take tokens as section names directly (best-effort)
    missing = []
    if not found and tokens:
        for t in tokens:
            # match closest available section containing token
            matched = [s for s in available_sections if t in s.lower()]
            if matched:
                for m in matched:
                    if m not in found:
                        found.append(m)
            else:
                # user requested a section not present
                missing.append(t)

    return found, missing

def enumerate_available_sections(data_dir: str) -> List[str]:
    """
    Scan uploaded files and attempt to detect section headings.
    Returns a de-duplicated list of candidate section names found.
    Strategies used:
      - For DOCX: scan paragraphs with heading style.
      - For PDF: attempt naive detection by searching for uppercase / common headings across text.
    This is a best-effort helper to give the parser names to match against.
    """
    sections = set()
    # DOCX scanning
    if Document is not None:
        for fname in os.listdir(data_dir):
            if fname.lower().endswith(".docx"):
                try:
                    doc = Document(os.path.join(data_dir, fname))
                    for para in doc.paragraphs:
                        style = getattr(para, "style", None)
                        text = para.text.strip()
                        if not text:
                            continue
                        # Heuristic: treat headings by style or uppercase short lines as headings
                        is_heading = False
                        try:
                            if style and hasattr(style, "name") and style.name.lower().startswith("heading"):
                                is_heading = True
                        except Exception:
                            pass
                        if len(text) < 120 and (is_heading or text.isupper() or text.istitle()):
                            sections.add(text)
                except Exception:
                    continue

    # PDF scanning
    if fitz is not None:
        for fname in os.listdir(data_dir):
            if fname.lower().endswith(".pdf"):
                path = os.path.join(data_dir, fname)
                try:
                    doc = fitz.open(path)
                    # sample first N pages to find headings
                    for page in doc:
                        blocks = page.get_text("blocks")
                        for b in blocks:
                            text = (b[4] or "").strip()
                            if not text:
                                continue
                            # heuristics: short lines, uppercase or titlecase likely headings
                            if 0 < len(text) < 120 and (text.isupper() or text.istitle()):
                                sections.add(text)
                except Exception:
                    continue

    # Add some common default section names
    defaults = ["Introduction", "Abstract", "Clinical Findings", "Findings", "Results", "Discussion",
                "Tables", "Images", "Figures", "Summary", "Conclusion"]
    for d in defaults:
        sections.add(d)

    return sorted(sections)

def extract_section_text_from_docx(filepath: str, section_names: List[str]) -> Dict[str, str]:
    """
    Return mapping section_name -> combined text found in docx file. Best-effort.
    """
    out = {s: "" for s in section_names}
    if Document is None:
        return out
    try:
        doc = Document(filepath)
    except Exception:
        return out

    current = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # If paragraph exactly matches a requested section, switch
        if text in section_names:
            current = text
            continue
        # If heading style suggests heading and matches any requested section
        try:
            if para.style and hasattr(para.style, "name") and para.style.name.lower().startswith("heading"):
                # check if heading matches one of requested sections using case-insensitive compare
                for s in section_names:
                    if text.lower() == s.lower():
                        current = s
                        break
        except Exception:
            pass
        if current:
            out[current] += text + "\n"
    return out

def extract_section_text_from_pdf(filepath: str, section_names: List[str]) -> Dict[str, str]:
    """
    Return mapping section_name -> combined text found in pdf file. This is naive: it looks for section name lines
    and then includes text until a next section is found.
    """
    out = {s: "" for s in section_names}
    if fitz is None:
        return out
    try:
        doc = fitz.open(filepath)
    except Exception:
        return out

    # Build full text with page separators
    full = []
    for page in doc:
        try:
            full.append(page.get_text())
        except Exception:
            full.append("")

    combined = "\n".join(full)
    lines = [l.strip() for l in combined.splitlines() if l.strip()]

    current = None
    for line in lines:
        # check if line matches one of requested sections
        for s in section_names:
            if line.lower() == s.lower() or line.lower() in s.lower() or s.lower() in line.lower():
                current = s
                break
        else:
            # no break: not a heading line
            if current:
                out[current] += line + "\n"
    return out

def collect_sections_from_data(data_dir: str, requested_sections: List[str]) -> Dict[str, str]:
    """
    Iterate files in data_dir and gather text for requested_sections from both docx and pdfs.
    Returns mapping: section -> combined text (from all docs).
    """
    collected = {s: "" for s in requested_sections}
    for fname in os.listdir(data_dir):
        path = os.path.join(data_dir, fname)
        try:
            if fname.lower().endswith(".docx"):
                res = extract_section_text_from_docx(path, requested_sections)
            elif fname.lower().endswith(".pdf"):
                res = extract_section_text_from_pdf(path, requested_sections)
            else:
                continue
            for k, v in res.items():
                if v:
                    collected[k] += v + "\n"
        except Exception:
            continue
    return collected


def build_pdf_in_memory(sections_text: dict) -> bytes:
    """
    Robust ReportLab builder: ensures all titles/bodies are strings,
    avoids string+int concat errors, and logs any problematic entries.
    """
    # sanitize input into ordered list of (title, body) both as strings
    sanitized = []
    for k, v in (sections_text or {}).items():
        try:
            title = str(k) if k is not None else "(Untitled Section)"
        except Exception as e:
            logger.exception("Failed to stringify section title: %s", k)
            title = "(Untitled Section)"
        try:
            body_text = "" if v is None else str(v)
        except Exception as e:
            logger.exception("Failed to stringify body for section %s: %s", title, e)
            body_text = ""
        sanitized.append((title, body_text))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            rightMargin=20*mm, leftMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=14
    )

    story = []
    for idx, (title, body) in enumerate(sanitized):
        story.append(Paragraph(title, heading_style))
        story.append(Spacer(1, 4))

        if not body.strip():
            story.append(Paragraph("(No content found for this section)", body_style))
            story.append(Spacer(1, 8))
        else:
            # Normalize and split into paragraphs
            paragraphs = [p.strip() for p in body.replace("\r\n", "\n").split("\n\n") if p.strip()]
            for p in paragraphs:
                try:
                    p_wrapped = "\n".join(textwrap.wrap(str(p), 100))
                    story.append(Paragraph(p_wrapped.replace("\n", "<br/>"), body_style))
                    story.append(Spacer(1, 6))
                except Exception:
                    # fallback to safe insertion
                    logger.exception("Paragraph render failed for section '%s'. Inserting raw text.", title)
                    story.append(Paragraph(str(p), body_style))
                    story.append(Spacer(1, 6))

        if idx != len(sanitized) - 1:
            story.append(PageBreak())

    doc.build(story)
    buf.seek(0)
    return buf.read()
